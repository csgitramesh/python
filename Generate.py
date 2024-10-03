import json
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess
import os
import comtypes.client
 
# Generate pie charts from JSON data
def generate_pie_chart(json_data):
   #If json_data is a file path, load the JSON content from the file
   with open('result.json', 'r') as json_file:
        json_content = json.load(json_file)
   # Pass the JSON content to read_json
   data = pd.DataFrame(json_content)
 
   # Group the data by "Services" and "Experiment" columns
   grouped_data = data.groupby(['Application', 'Services', 'Experiment', 'Results']).size().reset_index(name='count')
 
   # Iterate over each unique service
   for service in data['Services'].unique():
       # Filter data for the current service
       service_data = grouped_data[grouped_data['Services'] == service]
 
       # Create labels and sizes for the pie chart
       labels = service_data['Experiment'].tolist()
       sizes = service_data['count'].tolist()
       results = service_data['Results'].tolist()
 
       # Create colors for the pie chart based on "PASS" or "FAIL" status
       colors = ['mediumseagreen' if result == 'PASS' else 'tomato' for result in results]
 
       # Create a pie chart
       fig, ax = plt.subplots()
       ax.pie(sizes, labels=labels, colors=colors, startangle=90, wedgeprops={'linewidth': 2, 'edgecolor': 'white'})
 
       # Add title with service name in bold and increased font size
       ax.set_title(f'{service}', weight='bold', fontsize=15)
 
       # Generate legend
       pass_patch = Patch(color='mediumseagreen', label='Pass')
       fail_patch = Patch(color='tomato', label='Fail')
       plt.legend(handles=[pass_patch, fail_patch], loc='best')
 
       # Save the pie chart
       plt.savefig(f'pie_chart_{service}.png')
       plt.close()
 
# Update the Word document with JSON data
def update_doc(json_data):
   with open('result.json', 'r') as json_file:
        json_content1 = json.load(json_file)
   data1 = pd.DataFrame(json_content1)
   with open('SRE CHAOS SUMMARY REPORT.json', 'r') as json_file:
        json_content = json.load(json_file)
       
   # Pass the JSON content to read_json
   data = pd.DataFrame(json_content)  
   data_dict = {row["Requirments"]: [row["Details"], row["Size"], row["Style"]] for index, row in data.iterrows()}
   # Load the template Word document
   doc = Document(r'C:\Users\\DELL\\Desktop\\Report-Generation\\SRE CHAOS SUMMARY REPORT.docx')
   
   # Function to update placeholders in the document
   def update_placeholders():
       for table in doc.tables:
           for row in table.rows:
               for cell in row.cells:
                   for paragraph in cell.paragraphs:
                       for key, value in data_dict.items():
                           if key in paragraph.text:
                               # Replace the placeholder with the corresponding value
                               paragraph.text = paragraph.text.replace(key, value[0])
                               # Apply formatting to the replaced text
                               for run in paragraph.runs:
                                   run.font.size = Pt(value[1])
                                   run.font.bold = bool(value[2])
                                   run.font.name = 'Calibri (Body)'
   # Update placeholders in the document
   update_placeholders()
 
   # Find the "Chaos Test Results" section
   for i, paragraph in enumerate(doc.paragraphs):
       if 'Chaos Test Results' in paragraph.text:
           # Increase font font size for chaos test result section
           paragraph.style.font.size = Pt(14)
           # Iterate over the pie chart images and add them to the document
           for service in data1['Services'].unique():
               pie_chart_paragraph = doc.add_paragraph()
               pie_chart_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
               pie_chart_paragraph.add_run().add_picture(f'pie_chart_{service}.png', width=Inches(5))
           # Align the paragraph containing the pie charts to the left
           doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
           break
   # Add a reference section
   p = doc.add_paragraph()
   ref_run = p.add_run('Reference:')
   ref_run.bold = True
   ref_run.font.size = Pt(14) # Increase font size for reference section
 
   # Add a clickable link with blue color
   p = doc.add_paragraph()
   ref_link = 'https://zap.delta.com/ccoe/docs/operations_process_and_tooling/chaos-engineering/overview/'
   ref_run = p.add_run(ref_link)
   ref_run.italic = True
   ref_run.font.size = Pt(14) # Increase font size for the reference section
   ref_run.font.color.rgb = RGBColor(0, 0, 255) # Blue color
   
   # Save the modified document
   doc.save('SRE-CHAOS-SUMMARY-REPORT-APPNAME.docx')
 
   # Convert the Word document to PDF
   word_path = "SRE-CHAOS-SUMMARY-REPORT-APPNAME.docx"
   pdf_path = "SRE-CHAOS-SUMMARY-REPORT-APPNAME.pdf"
   
   
   # Create a COM object for Word application
   word = comtypes.client.CreateObject("Word.Application")
   docx_path = os.path.abspath(word_path)
   pdf_path = os.path.abspath(pdf_path)
   # PDF format code
   pdf_format = 17
   # Hide the Word application window
   word.Visible = False
   # Open the Word document
   in_file = word.Documents.Open(docx_path)
   # Save the Word document as PDF
   in_file.SaveAs(pdf_path, FileFormat=pdf_format)
   # Close the Word document
   in_file.Close()
   # Quit the Word application
   word.Quit()
 
# Read the Json files
json_detailed_report = "C:\\Data\\Chaos\\aws-lambda-chaos-library\\result.json"
json_summary_report = "C:\\Data\\Chaos\\aws-lambda-chaos-library\\SRE CHAOS SUMMARY REPORT.json"
 
# Generate pie charts
generate_pie_chart(json_detailed_report)
 
#Update the Word document
update_doc(json_summary_report)
